#!/usr/bin/env python3

from pathlib import Path
import pandas as pd

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("output/resubmission_pooled/20260819_all_sites_with_michigan")
DOCX_PATH = OUT_DIR / "leave_one_site_out_condensed_formal_table.docx"
CSV_PATH = OUT_DIR / "leave_one_site_out_condensed_formal_table.csv"
XLSX_PATH = OUT_DIR / "leave_one_site_out_condensed_formal_table.xlsx"


def polish_text(df):
    replacements = {
        "NO2": "NO\u2082",
        "PM25": "PM\u2082.\u2085",
        "PM2.5": "PM\u2082.\u2085",
        "ug/m3": "\u03bcg/m\u00b3",
        "I2": "I\u00b2",
    }
    out = df.copy()
    for col in out.columns:
        if out[col].dtype == object:
            for old, new in replacements.items():
                out[col] = out[col].str.replace(old, new, regex=False)
    return out


def read_csv(name):
    return polish_text(pd.read_csv(OUT_DIR / name, dtype=str).fillna(""))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="B7B7B7", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_column_widths(table, widths_dxa):
    set_table_width(table, sum(widths_dxa))
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def format_cell(cell, header=False, align="left"):
    set_cell_margins(cell)
    set_cell_border(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        set_cell_shading(cell, "EDEDED")
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(7.2)
            if header:
                run.bold = True


def add_dataframe_table(doc, df, widths, numeric_cols=None):
    numeric_cols = set(numeric_cols or [])
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, col in enumerate(df.columns):
        header.cells[i].text = str(col)
        format_cell(header.cells[i], header=True, align="center")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
            format_cell(cells[i], align="center" if col in numeric_cols else "left")
    set_column_widths(table, widths)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Note. ")
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(8)
    rest = p.add_run(text)
    rest.font.name = "Arial"
    rest.font.size = Pt(8)


def clean_primary(df):
    out = df.copy()
    out["Analysis"] = out["Analysis"].replace({
        "Primary mortality": "Mortality",
        "Primary ventilator-free days": "Ventilator-free days",
        "Cox mortality sensitivity": "Cox mortality sensitivity",
    })
    return out[
        [
            "Analysis",
            "Model",
            "Contrast",
            "Estimand",
            "All-site estimate (95% CI)",
            "Leave-one-site-out estimate range",
            "All-site I2, %",
            "Leave-one-site-out I2 range, %",
            "Maximum absolute change",
            "Leave-one estimates p<0.05, n",
        ]
    ].rename(columns={
        "All-site estimate (95% CI)": "All-site estimate (95% CI)",
        "Leave-one-site-out estimate range": "LOSO estimate range",
        "All-site I2, %": "All-site I\u00b2, %",
        "Leave-one-site-out I2 range, %": "LOSO I\u00b2 range, %",
        "Maximum absolute change": "Max change",
        "Leave-one estimates p<0.05, n": "LOSO p<0.05, n",
    })


def clean_fine_gray(df):
    out = df.copy()
    out = out[
        [
            "Outcome",
            "Model",
            "Contrast",
            "Estimand",
            "All-site estimate (95% CI)",
            "Leave-one-site-out estimate range",
            "All-site I2, %",
            "Leave-one-site-out I2 range, %",
            "Maximum absolute change",
            "Leave-one estimates p<0.05, n",
        ]
    ].rename(columns={
        "Leave-one-site-out estimate range": "LOSO estimate range",
        "All-site I2, %": "All-site I\u00b2, %",
        "Leave-one-site-out I2 range, %": "LOSO I\u00b2 range, %",
        "Maximum absolute change": "Max change",
        "Leave-one estimates p<0.05, n": "LOSO p<0.05, n",
    })
    return out


def build_outputs():
    primary = clean_primary(read_csv("leave_one_site_out_primary_summary_table.csv"))
    fine_gray = clean_fine_gray(read_csv("leave_one_site_out_fine_gray_summary_table.csv"))

    combined = pd.concat(
        [
            primary.assign(Table="Primary outcomes and Cox sensitivity"),
            fine_gray.assign(Table="Adjusted Fine-Gray secondary outcomes"),
        ],
        ignore_index=True,
        sort=False,
    )
    combined.to_csv(CSV_PATH, index=False)
    with pd.ExcelWriter(XLSX_PATH, engine="openpyxl") as writer:
        primary.to_excel(writer, sheet_name="Primary", index=False)
        fine_gray.to_excel(writer, sheet_name="Fine-Gray", index=False)
        combined.to_excel(writer, sheet_name="Combined", index=False)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9)
    for style_name in ("Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Condensed Leave-One-Site-Out Sensitivity Analysis")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(15)

    add_caption(doc, "Supplemental Table X. Leave-one-site-out analysis for primary outcomes")
    add_dataframe_table(
        doc,
        primary,
        [1250, 1450, 1250, 900, 1300, 1100, 800, 1050, 800, 900],
        numeric_cols={"All-site I\u00b2, %", "LOSO I\u00b2 range, %", "Max change", "LOSO p<0.05, n"},
    )
    add_note(
        doc,
        "LOSO = leave-one-site-out. Each row shows the all-site pooled estimate and the range of pooled estimates after sequentially excluding each site. I\u00b2 summarizes between-site heterogeneity.",
    )

    doc.add_page_break()
    add_caption(doc, "Supplemental Table X. Leave-one-site-out analysis for adjusted Fine-Gray secondary outcomes")
    add_dataframe_table(
        doc,
        fine_gray,
        [1450, 1450, 1250, 800, 1300, 1100, 800, 1050, 800, 900],
        numeric_cols={"All-site I\u00b2, %", "LOSO I\u00b2 range, %", "Max change", "LOSO p<0.05, n"},
    )
    add_note(
        doc,
        "Fine-Gray estimates are subdistribution hazard ratios for the secondary competing-risk outcomes among patients receiving invasive mechanical ventilation.",
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_outputs()
    print(DOCX_PATH)
