#!/usr/bin/env python3

from pathlib import Path
import math
import pandas as pd

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("output/resubmission_pooled/20260819_all_sites_with_michigan")
DOCX_PATH = OUT_DIR / "supplement_model_diagnostics_word_ready_tables.docx"


def read_csv(name):
    return polish_text(pd.read_csv(OUT_DIR / name, dtype=str).fillna(""))


def polish_text(df):
    replacements = {
        "NO2": "NO\u2082",
        "PM25": "PM\u2082.\u2085",
        "PM2.5": "PM\u2082.\u2085",
        "O3": "O\u2083",
        "ug/m3": "\u03bcg/m\u00b3",
        ">=24": "\u226524",
        "<=24": "\u226424",
    }
    out = df.copy()
    for col in out.columns:
        if out[col].dtype == object:
            for old, new in replacements.items():
                out[col] = out[col].str.replace(old, new, regex=False)
    return out


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="B7B7B7", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_column_widths(table, widths_dxa):
    set_table_width(table, sum(widths_dxa))
    grid = table._tbl.tblGrid
    if grid is None:
      grid = OxmlElement("w:tblGrid")
      table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def apply_doc_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9)

    for style_name in ("Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True

    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(11)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Note. ")
    run.italic = True
    run.font.size = Pt(8)
    run.font.name = "Arial"
    rest = p.add_run(text)
    rest.font.size = Pt(8)
    rest.font.name = "Arial"


def format_cell(cell, header=False, align="left"):
    set_cell_margins(cell)
    set_cell_border(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        set_cell_shading(cell, "EDEDED")
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(7.4)
            if header:
                run.bold = True


def add_dataframe_table(doc, df, widths, numeric_cols=None):
    numeric_cols = set(numeric_cols or [])
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, col in enumerate(df.columns):
        hdr.cells[i].text = str(col)
        format_cell(hdr.cells[i], header=True, align="center")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
            align = "center" if col in numeric_cols else "left"
            format_cell(cells[i], header=False, align=align)
    set_column_widths(table, widths)
    return table


def keep_cols(df, cols):
    return df.loc[:, cols].copy()


def build_doc():
    doc = Document()
    apply_doc_styles(doc)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Supplemental Model Diagnostics Tables")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("Pooled across CLIF sites; formatted for direct copy/paste into Word.")
    r.font.name = "Arial"
    r.font.size = Pt(9)

    vif = read_csv("supplement_vif_diagnostics_table.csv")
    vif = keep_cols(
        vif,
        [
            "Covariate",
            "Term type",
            "Model df",
            "Sites, n",
            "Median adjusted GVIF",
            "75th percentile adjusted GVIF",
            "Maximum adjusted GVIF",
            "Rows with adjusted GVIF >2.5",
            "Rows with adjusted GVIF >5",
            "Rows with adjusted GVIF >10",
        ],
    )
    add_caption(doc, "Supplemental Table X. Variance inflation diagnostics for adjusted competing-risk model covariates")
    add_dataframe_table(
        doc,
        vif,
        [2100, 1050, 700, 700, 1150, 1300, 1200, 1100, 1000, 1000],
        numeric_cols=set(vif.columns) - {"Covariate", "Term type"},
    )
    add_note(
        doc,
        "GVIF = generalized variance inflation factor. Adjusted GVIF is reported for multi-degree-of-freedom terms and is comparable with ordinary VIF for one-degree terms. No covariate had adjusted GVIF >5.",
    )

    doc.add_page_break()

    fg_ph = read_csv("supplement_fine_gray_ph_diagnostics_table.csv")
    fg_ph = keep_cols(
        fg_ph,
        [
            "Outcome",
            "Model",
            "Term",
            "Term type",
            "Sites, n",
            "Events, n",
            "Median site p",
            "Minimum site p",
            "Fisher combined p",
            "Sites with p<0.05, n",
        ],
    )
    add_caption(doc, "Supplemental Table X. Fine-Gray proportional hazards diagnostics")
    add_dataframe_table(
        doc,
        fg_ph,
        [1750, 1450, 1450, 800, 650, 800, 900, 900, 1050, 1050],
        numeric_cols={
            "Sites, n",
            "Events, n",
            "Median site p",
            "Minimum site p",
            "Fisher combined p",
            "Sites with p<0.05, n",
        },
    )
    add_note(
        doc,
        "Diagnostics summarize site-level cox.zph-style tests applied to the Fine-Gray model representation. Rows emphasize global tests and exposure terms.",
    )

    doc.add_page_break()

    fg_time = read_csv("supplement_fine_gray_time_interaction_diagnostics_table.csv")
    fg_time = keep_cols(
        fg_time,
        [
            "Outcome",
            "Model",
            "Exposure",
            "Sites, n",
            "Events, n",
            "Median exposure x log(time) coefficient",
            "Median site p",
            "Minimum site p",
            "Fisher combined p",
            "Sites with p<0.05, n",
        ],
    )
    add_caption(doc, "Supplemental Table X. Fine-Gray exposure-by-time interaction diagnostics")
    add_dataframe_table(
        doc,
        fg_time,
        [1750, 1450, 1300, 650, 800, 1400, 850, 850, 1000, 1050],
        numeric_cols=set(fg_time.columns) - {"Outcome", "Model", "Exposure"},
    )
    add_note(
        doc,
        "Time-interaction terms evaluate whether the exposure association varies with log(time). Small p values indicate evidence against a constant subdistribution hazard association over follow-up.",
    )

    doc.add_page_break()

    qp = read_csv("supplement_quasipoisson_dispersion_table.csv")
    qp = keep_cols(
        qp,
        [
            "Outcome",
            "Model",
            "Sites, n",
            "Participants, n",
            "Weighted mean outcome",
            "Weighted variance-to-mean ratio",
            "Median quasi-Poisson dispersion",
            "Dispersion range",
            "Median SE inflation vs Poisson",
            "Sites with Poisson overdispersion p<0.05, n",
        ],
    )
    add_caption(doc, "Supplemental Table X. Quasi-Poisson dispersion diagnostics for ventilator-free days and IMV duration")
    add_dataframe_table(
        doc,
        qp,
        [1850, 1350, 650, 900, 950, 1200, 1200, 1000, 1150, 1200],
        numeric_cols=set(qp.columns) - {"Outcome", "Model", "Dispersion range"},
    )
    add_note(
        doc,
        "Dispersion values >1 and quasi-Poisson/Poisson standard error ratios >1 indicate overdispersion relative to a Poisson variance assumption.",
    )

    pq = read_csv("supplement_poisson_vs_quasipoisson_table.csv")
    pq = keep_cols(
        pq,
        [
            "Outcome",
            "Model",
            "Term",
            "Sites, n",
            "Median quasi-Poisson dispersion",
            "Median quasi/Poisson SE ratio",
            "Median quasi-Poisson ratio of means",
            "Median Poisson ratio of means",
            "Sites with quasi-Poisson p<0.05, n",
            "Sites with Poisson p<0.05, n",
        ],
    )
    add_caption(doc, "Supplemental Table X. Comparison of quasi-Poisson and Poisson inference for exposure terms")
    add_dataframe_table(
        doc,
        pq,
        [1750, 1250, 1250, 650, 1050, 1050, 1150, 1100, 1100, 1100],
        numeric_cols=set(pq.columns) - {"Outcome", "Model", "Term"},
    )
    add_note(
        doc,
        "Point estimates are unchanged between Poisson and quasi-Poisson models with the same mean structure; quasi-Poisson inflates standard errors to account for overdispersion.",
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
